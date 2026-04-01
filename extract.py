import docx
import os

doc_path = r"d:\Kartik\OS_MiniProj\CPU_Scheduling_Simulator_Design_Requirement_Document.docx"
prd_path = r"d:\Kartik\OS_MiniProj\PRD.md"

trd_content = """

---

# 🧾 Technical Requirement Document (TRD)

## 1. System Overview

The system is a **Python-based CPU Scheduling Simulator** with:

* 6 scheduling algorithms
* Interactive dashboard (Streamlit)
* File upload support (CSV/JSON)
* Visualization (Gantt + graphs)
* Comparison + insights engine

---

## 2. Tech Stack

### Core

* **Language:** Python 3.x
* **Framework:** Streamlit

### Libraries

* **Pandas** → data handling
* **Plotly** → charts & Gantt
* **NumPy (optional)** → calculations

---

## 3. System Architecture

### High-Level Flow

```
User Input → Validation → Scheduling Engine → Metrics Engine → Visualization → Comparison → Insights
```

---

## 4. Module Breakdown

### 4.1 Input Module

Handles:

* Manual input (UI form)
* File upload (CSV/JSON)

Validation:

* No negative values
* Unique PID
* Required fields present

---

### 4.2 Core Data Model

```python
class Process:
    pid: str
    arrival: int
    burst: int
    priority: int
    remaining: int
    start_time: int
    completion_time: int
```

---

### 4.3 Scheduling Engine

#### Algorithms (6)

* FCFS
* SJF (Non-preemptive)
* Priority (Non-preemptive)
* SRTF
* Priority (Preemptive)
* Round Robin

---

### 4.4 Output Contract (Standard for all algorithms)

```python
{
  "processes": [...],
  "gantt": [...],
  "metrics": {...}
}
```

---

### 4.5 Metrics Engine

Compute:

* Turnaround Time = CT - AT
* Waiting Time = TAT - BT
* Response Time = First Start - AT
* CPU Utilization
* Throughput
* Context Switches

---

### 4.6 Visualization Module

#### Gantt Chart

* Plotly timeline
* Color-coded processes
* Idle time support

#### Graphs

* Bar charts:

  * Avg WT
  * Avg TAT
  * CPU Utilization
  * Context Switches

---

### 4.7 Comparison Engine

* Runs all 6 algorithms on same dataset
* Outputs:

  * Comparison table
  * Graphs
  * Ranking

---

### 4.8 Insights Engine

Generates:

* Best algorithm
* Worst algorithm
* Starvation warning
* Fairness analysis
* Context switch impact

---

## 5. File Handling

### Supported Formats

* CSV
* JSON

### CSV Format

```
pid,arrival,burst,priority
P1,0,5,2
P2,1,3,1
```

---

## 6. UI Structure (Streamlit)

### Sidebar

* Algorithm selection
* File upload
* Quantum input
* Run button

### Main Tabs

* Input Preview
* Results
* Gantt Chart
* Comparison
* Insights

---

## 7. Key Constraints

* Same input must be used for all algorithms
* Tie-breaking must be consistent
* No negative times allowed
* Gantt chart must match computed results
* Round Robin must follow strict FIFO queue

---

## 8. Performance Considerations

* Handle up to ~100 processes smoothly
* Avoid nested unnecessary loops
* Use efficient sorting and queue structures

---

## 9. Error Handling

* Invalid file format
* Missing fields
* Duplicate PID
* Negative values
* Empty input

---

## 10. Testing Requirements

### Must test:

* Each algorithm correctness
* Edge cases (idle CPU, same arrival)
* Round Robin queue behavior
* File upload parsing
* Comparison accuracy

---

## 11. Deployment (Optional)

* Run locally via Streamlit
* Optional:

  * Deploy on Streamlit Cloud
"""

def extract_and_save():
    doc = docx.Document(doc_path)
    extracted_text = []
    for para in doc.paragraphs:
        if para.text.strip():
            extracted_text.append(para.text)
            
    # Also extract tables if any
    for table in doc.tables:
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            extracted_text.append(" | ".join(row_data))
            
    # Convert extracted text to a simplified markdown structure
    md_content = "# Product Requirements Document (PRD)\n\n"
    
    # Simple heuristic to make all caps lines headings, etc., or just output text
    for line in extracted_text:
        if line.isupper() and len(line) < 50:
            md_content += f"\n## {line}\n\n"
        else:
            md_content += f"{line}\n\n"
            
    md_content += trd_content
    
    with open(prd_path, 'w', encoding='utf-8') as f:
        f.write(md_content)
        
    print(f"Successfully generated {prd_path}")

if __name__ == "__main__":
    extract_and_save()
